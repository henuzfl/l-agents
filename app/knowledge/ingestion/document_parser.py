from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import PdfReadError

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}


class InvalidKnowledgeDocument(ValueError):
    pass


def extract_document_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        allowed = "、".join(sorted(SUPPORTED_EXTENSIONS))
        raise InvalidKnowledgeDocument(f"不支持该文件类型，可上传：{allowed}")
    try:
        if suffix in {".txt", ".md", ".markdown"}:
            text = content.decode("utf-8-sig")
        elif suffix == ".pdf":
            pages = PdfReader(BytesIO(content)).pages
            text = "\n".join(page.extract_text() or "" for page in pages)
        else:
            document = DocxDocument(BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    except (
        BadZipFile,
        KeyError,
        PackageNotFoundError,
        PdfReadError,
        UnicodeDecodeError,
        ValueError,
    ) as exc:
        raise InvalidKnowledgeDocument("无法解析文档，请确认文件未损坏且内容格式正确。") from exc
    cleaned = text.strip()
    if not cleaned:
        raise InvalidKnowledgeDocument("文档中没有可提取的文本。")
    return cleaned
