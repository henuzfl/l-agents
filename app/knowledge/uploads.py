from __future__ import annotations

import mimetypes
from collections.abc import Callable
from copy import deepcopy
from dataclasses import asdict, dataclass, field
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from threading import Lock
from typing import Literal
from uuid import uuid4
from zipfile import BadZipFile

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from llama_index.core import Document
from minio import Minio
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.core.config import Settings
from app.core.exceptions import KnowledgeConfigurationError, KnowledgeRetrievalError

from .documents import build_document_nodes
from .pdf_processing import process_pdf
from .registry import KnowledgeDocumentRegistry
from .store import KnowledgeStatus, LlamaIndexKnowledgeStore

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}
StageState = Literal["pending", "running", "completed", "failed"]
JobState = Literal["processing", "completed", "failed"]


class InvalidKnowledgeDocument(ValueError):
    pass


StoreFactory = Callable[[Settings], LlamaIndexKnowledgeStore]


@dataclass
class PipelineStage:
    state: StageState = "pending"
    detail: str = "等待处理"


@dataclass
class DocumentJob:
    task_id: str
    filename: str
    object_name: str
    status: JobState
    loading: PipelineStage
    chunking: PipelineStage
    embedding: PipelineStage
    chunk_count: int | None
    created_at: str
    updated_at: str
    error: str | None = None
    warnings: list[str] = field(default_factory=list)
    element_counts: dict[str, int] = field(default_factory=dict)

    def to_dict(self) -> dict[str, object]:
        return asdict(self)

    @classmethod
    def from_dict(cls, payload: dict[str, object]) -> DocumentJob:
        data = dict(payload)
        data.setdefault("warnings", [])
        data.setdefault("element_counts", {})
        data["loading"] = PipelineStage(**data["loading"])  # type: ignore[arg-type]
        data["chunking"] = PipelineStage(**data["chunking"])  # type: ignore[arg-type]
        data["embedding"] = PipelineStage(**data["embedding"])  # type: ignore[arg-type]
        return cls(**data)  # type: ignore[arg-type]


class MinioDocumentStorage:
    def __init__(self, settings: Settings, client: Minio | None = None) -> None:
        if not settings.minio_endpoint:
            raise KnowledgeConfigurationError("缺少 MinIO 配置：MINIO_ENDPOINT")
        if not settings.minio_access_key:
            raise KnowledgeConfigurationError("缺少 MinIO 配置：MINIO_ACCESS_KEY")
        if settings.minio_secret_key is None:
            raise KnowledgeConfigurationError("缺少 MinIO 配置：MINIO_SECRET_KEY")
        self._bucket = settings.minio_bucket
        self._client = client or Minio(
            settings.minio_endpoint,
            access_key=settings.minio_access_key,
            secret_key=settings.minio_secret_key.get_secret_value(),
            secure=settings.minio_secure,
        )

    @property
    def bucket(self) -> str:
        return self._bucket

    def upload(self, object_name: str, content: bytes, content_type: str) -> None:
        try:
            if not self._client.bucket_exists(self._bucket):
                self._client.make_bucket(self._bucket)
            self._client.put_object(
                self._bucket,
                object_name,
                BytesIO(content),
                length=len(content),
                content_type=content_type,
            )
        except Exception as exc:
            raise KnowledgeRetrievalError("原始文档上传 MinIO 失败。") from exc

    def download(self, object_name: str) -> bytes:
        response = None
        try:
            response = self._client.get_object(self._bucket, object_name)
            return response.read()
        except Exception as exc:
            raise KnowledgeRetrievalError("无法从 MinIO 读取原始文档。") from exc
        finally:
            if response is not None:
                response.close()
                response.release_conn()

    def delete(self, object_name: str) -> None:
        try:
            self._client.remove_object(self._bucket, object_name)
        except Exception as exc:
            raise KnowledgeRetrievalError("MinIO 原始文档删除失败。") from exc

    def delete_assets(self, object_name: str) -> None:
        asset_prefix = f"{object_name.rsplit('/', 1)[0]}/assets/"
        try:
            for item in self._client.list_objects(
                self._bucket,
                prefix=asset_prefix,
                recursive=True,
            ):
                self._client.remove_object(self._bucket, item.object_name)
        except Exception as exc:
            raise KnowledgeRetrievalError("MinIO 文档图片清理失败。") from exc


def extract_document_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        allowed = "、".join(sorted(SUPPORTED_EXTENSIONS))
        raise InvalidKnowledgeDocument(f"不支持该文件类型，可上传：{allowed}")
    try:
        if suffix in {".txt", ".md", ".markdown"}:
            text = content.decode("utf-8-sig")
        elif suffix == ".pdf":
            pages = PdfReader(BytesIO(content)).pages
            text = "\n".join(page.extract_text() or "" for page in pages)
        else:
            document = DocxDocument(BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    except (
        BadZipFile,
        KeyError,
        PackageNotFoundError,
        PdfReadError,
        UnicodeDecodeError,
        ValueError,
    ) as exc:
        raise InvalidKnowledgeDocument("无法解析文档，请确认文件未损坏且内容格式正确。") from exc
    cleaned = text.strip()
    if not cleaned:
        raise InvalidKnowledgeDocument("文档中没有可提取的文本。")
    return cleaned


class KnowledgeDocumentService:
    def __init__(
        self,
        settings: Settings,
        store_factory: StoreFactory = LlamaIndexKnowledgeStore,
        object_storage: MinioDocumentStorage | None = None,
        registry: KnowledgeDocumentRegistry | None = None,
    ) -> None:
        self._settings = settings
        self._store_factory = store_factory
        self._object_storage = object_storage
        self._registry = registry or KnowledgeDocumentRegistry(settings.knowledge_registry_path)
        self._jobs: dict[str, DocumentJob] = {}
        self._lock = Lock()
        self._recover_interrupted_jobs()

    @property
    def max_upload_bytes(self) -> int:
        return self._settings.knowledge_upload_max_bytes

    def _storage(self) -> MinioDocumentStorage:
        if self._object_storage is None:
            self._object_storage = MinioDocumentStorage(self._settings)
        return self._object_storage

    def _recover_interrupted_jobs(self) -> None:
        for payload in self._registry.list():
            job = DocumentJob.from_dict(payload)
            if job.status != "processing":
                continue
            job.status = "failed"
            job.error = "服务重启导致后台处理被中断，请重新处理。"
            if job.chunking.state == "running":
                job.chunking = PipelineStage("failed", "后台处理已中断")
            elif job.embedding.state == "running":
                job.embedding = PipelineStage("failed", "后台处理已中断")
            job.updated_at = self._now()
            self._registry.save(job.to_dict())

    @staticmethod
    def _now() -> str:
        return datetime.now(UTC).isoformat()

    def _update(self, task_id: str, **changes: object) -> None:
        with self._lock:
            job = self._jobs[task_id]
            for name, value in changes.items():
                setattr(job, name, value)
            job.updated_at = self._now()
            self._registry.save(job.to_dict())

    def start_upload(self, filename: str, content: bytes, content_type: str | None) -> DocumentJob:
        clean_filename = Path(filename).name
        if not clean_filename:
            raise InvalidKnowledgeDocument("文件名不能为空。")
        if Path(clean_filename).suffix.lower() not in SUPPORTED_EXTENSIONS:
            allowed = "、".join(sorted(SUPPORTED_EXTENSIONS))
            raise InvalidKnowledgeDocument(f"不支持该文件类型，可上传：{allowed}")
        if len(content) > self._settings.knowledge_upload_max_bytes:
            limit_mb = self._settings.knowledge_upload_max_bytes // (1024 * 1024)
            raise InvalidKnowledgeDocument(f"文件不能超过 {limit_mb} MB。")

        task_id = str(uuid4())
        date_path = datetime.now(UTC).strftime("%Y/%m/%d")
        object_name = f"{date_path}/{task_id}/{clean_filename}"
        now = self._now()
        job = DocumentJob(
            task_id=task_id,
            filename=clean_filename,
            object_name=object_name,
            status="processing",
            loading=PipelineStage("running", "正在上传原始文档到 MinIO"),
            chunking=PipelineStage(),
            embedding=PipelineStage(),
            chunk_count=None,
            created_at=now,
            updated_at=now,
        )
        with self._lock:
            self._jobs[task_id] = job
            self._registry.save(job.to_dict())
        detected_type = content_type or mimetypes.guess_type(clean_filename)[0]
        try:
            self._storage().upload(
                object_name,
                content,
                detected_type or "application/octet-stream",
            )
        except Exception as exc:
            self._update(
                task_id,
                status="failed",
                loading=PipelineStage("failed", "原始文档上传失败"),
                error=str(exc),
            )
            raise
        self._update(
            task_id,
            loading=PipelineStage("completed", f"已保存到 {self._storage().bucket}"),
            chunking=PipelineStage("pending", "后台任务等待开始"),
        )
        return self.get_job(task_id)

    def process(self, task_id: str) -> None:
        try:
            self._update(task_id, chunking=PipelineStage("running", "正在提取并切分文本"))
            job = self.get_job(task_id)
            content = self._storage().download(job.object_name)
            if Path(job.filename).suffix.lower() == ".pdf":
                pdf_result = process_pdf(
                    content,
                    job.filename,
                    job.object_name,
                    self._settings,
                    self._storage().upload,
                )
                nodes = pdf_result.nodes
                warnings = pdf_result.warnings
                element_counts = pdf_result.element_counts
            else:
                text = extract_document_text(job.filename, content)
                document = Document(
                    text=text,
                    metadata={
                        "source": job.filename,
                        "section": "上传文档",
                        "minio_bucket": self._storage().bucket,
                        "minio_object": job.object_name,
                    },
                )
                nodes = build_document_nodes(document)
                warnings = []
                element_counts = {"text": len(nodes)}
            self._update(
                task_id,
                chunk_count=len(nodes),
                warnings=warnings,
                element_counts=element_counts,
                chunking=PipelineStage("completed", f"已生成 {len(nodes)} 个文本块"),
                embedding=PipelineStage("running", "正在生成向量并写入索引"),
            )
            self._store_factory(self._settings).add_nodes(nodes)
            self._update(
                task_id,
                status="completed",
                embedding=PipelineStage("completed", "向量索引写入完成"),
            )
        except Exception as exc:
            with self._lock:
                job = self._jobs[task_id]
                if job.chunking.state == "running":
                    job.chunking = PipelineStage("failed", "文档切分失败")
                elif job.embedding.state == "running":
                    job.embedding = PipelineStage("failed", "向量写入失败")
                job.status = "failed"
                job.error = str(exc)
                job.updated_at = self._now()
                self._registry.save(job.to_dict())

    def get_job(self, task_id: str) -> DocumentJob:
        with self._lock:
            job = self._jobs.get(task_id)
            if job is None:
                payload = self._registry.get(task_id)
                if payload is None:
                    raise KeyError(task_id)
                job = DocumentJob.from_dict(payload)
                self._jobs[task_id] = job
            return deepcopy(job)

    def list_jobs(self) -> list[dict[str, object]]:
        return self._registry.list()

    def download(self, task_id: str) -> tuple[str, str, bytes]:
        job = self.get_job(task_id)
        content_type = mimetypes.guess_type(job.filename)[0] or "application/octet-stream"
        return job.filename, content_type, self._storage().download(job.object_name)

    def download_asset(self, task_id: str, node_id: str) -> tuple[str, str, bytes]:
        job = self.get_job(task_id)
        if job.status != "completed":
            raise InvalidKnowledgeDocument("文档处理完成后才能查看图片。")
        chunk = self._store_factory(self._settings).get_document_chunk(job.object_name, node_id)
        asset_object = chunk.metadata.get("asset_object") if chunk is not None else None
        if not isinstance(asset_object, str) or not asset_object:
            raise KeyError(node_id)
        asset_prefix = f"{job.object_name.rsplit('/', 1)[0]}/assets/"
        if not asset_object.startswith(asset_prefix):
            raise InvalidKnowledgeDocument("分片图片不属于当前文档。")
        filename = Path(asset_object).name
        content_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
        return filename, content_type, self._storage().download(asset_object)

    def list_chunks(
        self,
        task_id: str,
        *,
        offset: int = 0,
        limit: int = 100,
        element_type: str | None = None,
        query: str | None = None,
    ) -> dict[str, object]:
        job = self.get_job(task_id)
        if job.status != "completed":
            raise InvalidKnowledgeDocument("文档处理完成后才能查看分片。")
        total, chunks = self._store_factory(self._settings).list_document_chunks(
            job.object_name,
            offset=offset,
            limit=limit,
            element_type=element_type,
            query=query,
        )
        return {
            "task_id": job.task_id,
            "filename": job.filename,
            "minio_bucket": self._storage().bucket,
            "minio_object": job.object_name,
            "download_url": f"/api/v1/knowledge/documents/{job.task_id}/download",
            "total": total,
            "offset": offset,
            "limit": limit,
            "chunks": [
                {
                    "position": offset + index,
                    "node_id": chunk.node_id,
                    "content": chunk.content,
                    "page_number": chunk.metadata.get("page_number"),
                    "section": chunk.metadata.get("section", "上传文档"),
                    "element_type": chunk.metadata.get("element_type", "text"),
                    "element_index": chunk.metadata.get("element_index"),
                    "element_part": chunk.metadata.get("element_part"),
                    "language": chunk.metadata.get("language"),
                    "minio_object": chunk.metadata.get("minio_object", job.object_name),
                    "asset_object": chunk.metadata.get("asset_object"),
                    "asset_url": (
                        f"/api/v1/knowledge/documents/{job.task_id}/chunks/"
                        f"{chunk.node_id}/asset"
                        if chunk.metadata.get("asset_object")
                        else None
                    ),
                }
                for index, chunk in enumerate(chunks, start=1)
            ],
        }

    def restart(self, task_id: str) -> DocumentJob:
        job = self.get_job(task_id)
        if job.status == "processing":
            raise InvalidKnowledgeDocument("文档正在处理中，不能重复执行。")
        self._store_factory(self._settings).delete_document(job.object_name)
        self._storage().delete_assets(job.object_name)
        self._update(
            task_id,
            status="processing",
            chunk_count=None,
            warnings=[],
            element_counts={},
            chunking=PipelineStage("pending", "后台任务等待开始"),
            embedding=PipelineStage(),
            error=None,
        )
        return self.get_job(task_id)

    def delete(self, task_id: str) -> None:
        job = self.get_job(task_id)
        if job.status == "processing":
            raise InvalidKnowledgeDocument("文档正在处理中，不能删除。")
        self._store_factory(self._settings).delete_document(job.object_name)
        self._storage().delete_assets(job.object_name)
        self._storage().delete(job.object_name)
        with self._lock:
            self._jobs.pop(task_id, None)
            self._registry.delete(task_id)

    def status(self) -> dict[str, int | str]:
        status: KnowledgeStatus = self._store_factory(self._settings).status()
        return asdict(status)
